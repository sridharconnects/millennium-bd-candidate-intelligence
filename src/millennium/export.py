"""Deliverable #2: parsed resume data as JSON and CSV.

Three shapes, because they answer different questions:

* `candidates.json`  -- full fidelity, every field with its evidence and status. This
  is the contract other systems integrate against.
* `candidates.csv`   -- one row per candidate, flattened for Excel. Unknown values are
  written as an empty cell and a companion `*_status` column says WHY: 'abstained'
  (we saw a claim we could not prove) reads very differently from 'missing' (the CV
  never said). Collapsing those two into a blank is the standard way this data gets
  quietly misread.
* `employment.csv` / `education.csv` / `skills.csv` -- long-form for pivoting.
"""
from __future__ import annotations

import csv
import json
from datetime import datetime, timezone
from pathlib import Path

from . import taxonomy as tx
from .config import SETTINGS
from .schema import CandidateProfile, Tracked


def _status(t: Tracked) -> str:
    if t.is_known:
        return t.validation_status
    return "abstained" if t.validation_status == "abstained" else "missing"


def _v(t: Tracked):
    return t.normalized_value if t.is_known and t.normalized_value is not None else (
        t.value if t.is_known else "")


def flat_row(p: CandidateProfile, include_pii: bool = True) -> dict:
    cur = p.current_role()
    highest = None
    order = {"phd": 5, "mba": 4, "professional": 4, "masters": 3, "bachelors": 2, "secondary": 1}
    for e in p.education:
        if e.degree_level and (highest is None or order.get(e.degree_level, 0) > order.get(highest, 0)):
            highest = e.degree_level

    row = {
        "candidate_id": p.candidate_id,
        "source_file": p.provenance.source_file if p.provenance else "",
        "is_synthetic": bool(p.provenance and p.provenance.is_synthetic),
        "headline": _v(p.headline), "headline_status": _status(p.headline),
        "location": _v(p.location_current), "location_status": _status(p.location_current),
        "country": p.geography.label if p.geography else "",
        "region": tx.display("region", p.geo_region.label, "") if p.geo_region else "",
        "region_confidence": round(p.geo_region.confidence, 3) if p.geo_region else "",
        "years_experience": _v(p.years_experience),
        "years_experience_status": _status(p.years_experience),
        "years_relevant": _v(p.years_relevant_experience),
        "seniority_level": p.seniority.label if p.seniority else "",
        "seniority_title": (tx.display("seniority", p.seniority.label, "")
                            if p.seniority and p.seniority.label.startswith("L") else ""),
        "current_employer": (cur.employer_canonical or cur.employer_raw.display(""))
                            if cur else "",
        "current_employer_tier": tx.display("tier", cur.employer_tier, "") if cur else "",
        "current_title": cur.title_raw.display("") if cur else "",
        "current_tenure_months": _v(p.current_tenure_months),
        "approach": p.quant_fundamental.label if p.quant_fundamental else "",
        "feeder_path": (tx.display("feeder", p.feeder_path.label)
                        if p.feeder_path else ""),
        "strategies": "; ".join(tx.display("strategy", c.label) for c in p.strategies),
        "strategies_confidence": "; ".join(f"{c.confidence:.2f}" for c in p.strategies),
        "sectors": "; ".join(tx.display("sector", c.label) for c in p.sectors),
        "skills_core": "; ".join(s.canonical for s in p.skills if s.depth == "core"),
        "skills_all": "; ".join(s.canonical for s in p.skills),
        "certifications": "; ".join(
            f"{tx.display("certification", c.canonical)}"
            + (f" ({c.status})" if c.status else "") for c in p.certifications if c.canonical),
        "languages": "; ".join(
            f"{l.language}" + (f" ({l.proficiency})" if l.proficiency else "")
            for l in p.languages),
        "highest_degree": highest or "",
        "institutions": "; ".join(e.institution.display("") for e in p.education
                                  if e.institution.is_known),
        "n_roles": len(p.employment),
        "n_employment_gaps": len(p.employment_gaps),
        "completeness": p.quality.completeness,
        "evidence_coverage": p.quality.evidence_coverage,
        "extraction_quality": p.quality.extraction_quality,
        "abstentions": p.quality.abstention_count,
        "needs_human_review": p.quality.needs_human_review,
        "review_reasons": " | ".join(p.quality.review_reasons),
        "validation_flags": " | ".join(p.quality.validation_flags),
        "injection_flags": "; ".join(p.provenance.injection_flags) if p.provenance else "",
        "near_duplicate_of": "; ".join(p.provenance.near_duplicate_of) if p.provenance else "",
        "llm_model": p.provenance.llm_model if p.provenance else "",
        "cost_usd": p.provenance.cost_usd if p.provenance else 0.0,
        "schema_version": p.provenance.schema_version if p.provenance else "",
    }
    if include_pii:
        row |= {
            "full_name": _v(p.sensitive.full_name),
            "full_name_status": _status(p.sensitive.full_name),
            "email": _v(p.sensitive.email), "email_status": _status(p.sensitive.email),
            "phone": _v(p.sensitive.phone), "phone_status": _status(p.sensitive.phone),
        }
    return row


def _write_csv(path: Path, rows: list[dict]) -> Path:
    if not rows:
        path.write_text("")
        return path
    keys: list[str] = []
    for r in rows:
        for k in r:
            if k not in keys:
                keys.append(k)
    with path.open("w", newline="", encoding="utf8") as f:
        wr = csv.DictWriter(f, fieldnames=keys, extrasaction="ignore")
        wr.writeheader()
        wr.writerows(rows)
    return path


def export_all(profiles: list[CandidateProfile], out_dir: Path | None = None,
               manifest: dict | None = None, include_pii: bool = True) -> dict[str, Path]:
    out = Path(out_dir or SETTINGS.paths.exports)
    out.mkdir(parents=True, exist_ok=True)
    written: dict[str, Path] = {}

    payload = {
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "schema_version": SETTINGS.schema_version,
        "taxonomy_version": tx.TAXONOMY_VERSION,
        "count": len(profiles),
        "manifest": manifest or {},
        "candidates": [json.loads(p.model_dump_json(exclude={"raw_text"})) for p in profiles],
    }
    pj = out / "candidates.json"
    pj.write_text(json.dumps(payload, indent=1, ensure_ascii=False), encoding="utf8")
    written["candidates.json"] = pj

    written["candidates.csv"] = _write_csv(
        out / "candidates.csv", [flat_row(p, include_pii) for p in profiles])

    emp_rows = []
    for p in profiles:
        for i, e in enumerate(p.employment):
            emp_rows.append({
                "candidate_id": p.candidate_id, "seq": i,
                "employer_raw": e.employer_raw.display(""),
                "employer_canonical": e.employer_canonical or "",
                "employer_tier": tx.display("tier", e.employer_tier, ""),
                "title": e.title_raw.display(""),
                "seniority_level": e.seniority_level or "",
                "location": e.location.display(""),
                "start": e.dates.start.normalized_value or "",
                "end": e.dates.end.normalized_value or "",
                "start_status": _status(e.dates.start), "end_status": _status(e.dates.end),
                "duration_months": _v(e.dates.duration_months),
                "is_current": e.dates.is_current, "is_internship": e.is_internship,
                "n_highlights": len(e.highlights),
                "evidence_page": (e.employer_raw.evidence[0].page
                                  if e.employer_raw.evidence else ""),
                "evidence_char_start": (e.employer_raw.evidence[0].char_start
                                        if e.employer_raw.evidence else ""),
            })
    written["employment.csv"] = _write_csv(out / "employment.csv", emp_rows)

    edu_rows = []
    for p in profiles:
        for e in p.education:
            edu_rows.append({
                "candidate_id": p.candidate_id,
                "institution": e.institution.display(""), "degree": e.degree_raw.display(""),
                "degree_level": e.degree_level or "", "field": e.field_of_study.display(""),
                "graduation_year": _v(e.graduation_year), "gpa": e.gpa_raw.display(""),
                "location": e.location.display(""), "honors": "; ".join(e.honors),
                "institution_status": _status(e.institution),
            })
    written["education.csv"] = _write_csv(out / "education.csv", edu_rows)

    skill_rows = [{"candidate_id": p.candidate_id, "skill": s.canonical,
                   "category": s.category, "depth": s.depth,
                   "surface_forms": "; ".join(s.surface_forms),
                   "n_evidence": len(s.evidence)}
                  for p in profiles for s in p.skills]
    written["skills.csv"] = _write_csv(out / "skills.csv", skill_rows)

    _CTX_PAD = 80

    def _context(p: CandidateProfile, ev) -> str:
        """A readable excerpt around the span, for a reviewer scanning the CSV without
        the app open. Built here from `raw_text`, not stored on `Evidence` itself --
        `Evidence.snippet` is defined to be the exact matched text (see classification.
        _ev and validate.verify_span), and conflating the two previously let a genuine
        span-integrity bug (padded text mislabelled as an exact match) go unnoticed."""
        if not p.raw_text:
            return ev.snippet[:200]
        lo = max(0, ev.char_start - _CTX_PAD)
        hi = min(len(p.raw_text), ev.char_end + _CTX_PAD)
        return p.raw_text[lo:hi].replace("\n", " ")[:260]

    ev_rows = [{"candidate_id": p.candidate_id, "doc_id": ev.doc_id, "page": ev.page or "",
                "char_start": ev.char_start, "char_end": ev.char_end,
                "match_kind": ev.match_kind, "match_score": ev.match_score,
                "snippet": ev.snippet[:200].replace("\n", " "),
                "context": _context(p, ev)}
               for p in profiles for ev in p.all_evidence()]
    written["evidence.csv"] = _write_csv(out / "evidence.csv", ev_rows)
    return written


def _slug(p: CandidateProfile, blind: bool = False) -> str:
    raw = p.display_name(blind).lower()
    keep = "".join(c if c.isalnum() else "-" for c in raw).strip("-")
    return (keep or p.candidate_id[:8])[:48]


def profile_filename(p: CandidateProfile, ext: str, blind: bool = False) -> str:
    return f"{_slug(p, blind)}.{ext.lstrip('.')}" if ext else _slug(p, blind)


def profile_json_bytes(p: CandidateProfile, include_pii: bool = True) -> bytes:
    payload = json.loads(p.model_dump_json(exclude={"raw_text"}))
    if not include_pii:
        payload.pop("sensitive", None)
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf8")


def profile_csv_bytes(p: CandidateProfile, include_pii: bool = True) -> bytes:
    row = flat_row(p, include_pii=include_pii)
    from io import StringIO
    buf = StringIO()
    wr = csv.DictWriter(buf, fieldnames=list(row.keys()))
    wr.writeheader()
    wr.writerow(row)
    return buf.getvalue().encode("utf8")


def _profile_lines(p: CandidateProfile, include_pii: bool) -> list[tuple[str, str]]:
    """Plain-text sections for PDF / Word. (heading, body)."""
    name = p.display_name(not include_pii)
    cur = p.current_role()
    lines: list[tuple[str, str]] = [
        ("Name", name),
        ("Headline", p.headline.display("—")),
        ("Location", p.location_current.display("—")),
        ("Years experience", p.years_experience.display("—")),
        ("Seniority", (f"{p.seniority.label} · {tx.display('seniority', p.seniority.label)}"
                       if p.seniority else "—")),
        ("Current role", (f"{cur.title_raw.display('—')} · "
                          f"{cur.employer_canonical or cur.employer_raw.display('—')}"
                          if cur else "—")),
        ("Strategies", ", ".join(tx.display("strategy", c.label) for c in p.strategies) or "—"),
        ("Sectors", ", ".join(tx.display("sector", c.label) for c in p.sectors) or "—"),
        ("Skills", ", ".join(s.canonical for s in p.skills) or "—"),
        ("Languages", ", ".join(
            f"{l.language}" + (f" ({l.proficiency})" if l.proficiency else "")
            for l in p.languages) or "—"),
        ("Certifications", ", ".join(
            tx.display("certification", c.canonical) for c in p.certifications if c.canonical) or "—"),
    ]
    if include_pii:
        lines.insert(2, ("Email", p.sensitive.email.display("—")))
        lines.insert(3, ("Phone", p.sensitive.phone.display("—")))
    emp = []
    for e in p.employment:
        dates = f"{e.dates.start.normalized_value or '?'} → {e.dates.end.normalized_value or '?'}"
        emp.append(f"{e.title_raw.display('—')} · "
                   f"{e.employer_canonical or e.employer_raw.display('—')} ({dates})")
        for h in e.highlights[:4]:
            emp.append(f"  – {h.value}")
    lines.append(("Employment", "\n".join(emp) if emp else "—"))
    edu = []
    for e in p.education:
        edu.append(" · ".join(x for x in [
            e.institution.display(""), e.degree_raw.display(""),
            e.field_of_study.display(""), str(e.graduation_year.display("")),
        ] if x and x != "—"))
    lines.append(("Education", "\n".join(edu) if edu else "—"))
    return lines


def profile_docx_bytes(p: CandidateProfile, include_pii: bool = True) -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = doc.add_heading(p.display_name(not include_pii), level=1)
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    sub = doc.add_paragraph(p.headline.display(""))
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    for heading, body in _profile_lines(p, include_pii):
        h = doc.add_heading(heading, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
        para = doc.add_paragraph(str(body))
        para.paragraph_format.space_after = Pt(8)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def profile_pdf_bytes(p: CandidateProfile, include_pii: bool = True) -> bytes:
    import pymupdf as fitz

    def wrap(text: str, width: int = 92) -> list[str]:
        out: list[str] = []
        for para in str(text).split("\n"):
            line = ""
            for word in para.split() or [""]:
                trial = (line + " " + word).strip()
                if len(trial) > width and line:
                    out.append(line)
                    line = word
                else:
                    line = trial
            out.append(line)
        return out or [""]

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y, margin = 56.0, 48.0

    def write(text: str, size: float, color=(0.06, 0.09, 0.16), bold: bool = False) -> None:
        nonlocal y, page
        font = "hebo" if bold else "helv"
        for line in wrap(text):
            if y > 800:
                page = doc.new_page(width=595, height=842)
                y = 56.0
            page.insert_text((margin, y), line, fontsize=size, fontname=font, color=color)
            y += size + 5

    write(p.display_name(not include_pii), 18, bold=True)
    write(p.headline.display(""), 11, color=(0.39, 0.45, 0.55))
    y += 8
    for heading, body in _profile_lines(p, include_pii):
        write(heading.upper(), 9, color=(0.06, 0.46, 0.43), bold=True)
        write(str(body), 10)
        y += 8
    data = doc.tobytes()
    doc.close()
    return data
