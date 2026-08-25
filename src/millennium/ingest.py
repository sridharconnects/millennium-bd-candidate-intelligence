"""Document ingestion: bytes -> clean, correctly-ordered text with page offsets.

Nothing here is speculative. Each repair below exists because a specific file in the
supplied corpus is broken in a specific way, verified by `scripts/inventory.py`:

* Omar El-Hassan 202405.pdf  -- two-column CV. Naive `page.get_text()` interleaves the
  right-hand skills sidebar into the middle of the work-experience bullets, so the
  contact line lands inside a job description. Block x0 values are cleanly bimodal
  (42-66 vs 418-428), so we cluster blocks into columns and read column-major.
  The same file has Type1 subsetting damage: U+019F is a 'ti' ligature and U+019E is
  'tf', which is why the raw text says 'QuanƟtaƟve' and 'Porƞolio'.
* Viktor Sharat.docx -- merged table cells. `row.cells` yields the same underlying
  <w:tc> element once per grid column, which duplicates every achievement four times.
  We de-duplicate on the identity of the underlying XML element.
* Michael Rodriguez, CFA.docx / Zara Al-Rashid.docx -- content lives in tables that
  python-docx's `paragraphs` property skips entirely, and appending tables afterwards
  destroys reading order (Michael's EDUCATION heading ends up after INTERESTS).
  We walk the body element children in true document order instead.
* Priya Nakamura ... .docx -- carries a 'RED LANE TALENT MANAGEMENT' agency watermark
  in a header. Headers are extracted separately and tagged, never mixed into the body,
  but they are retained because agency provenance is genuinely useful to a recruiter.
"""
from __future__ import annotations

import hashlib
import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.oxml.ns import qn

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# --------------------------------------------------------------------------- repair
# Ligature and Type1-subsetting damage. Order matters: longest first.
LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "ﬅ": "st", "ﬆ": "st",
    "Ɵ": "ti",   # Ɵ  -> 'ti'   (QuanƟtaƟve -> Quantitative)
    "ƞ": "tf",   # ƞ  -> 'tf'   (Porƞolio   -> Portfolio)
    "Ŧ": "TI", "ŧ": "ti",
    "—": "-", "–": "-", "―": "-", "‒": "-",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    " ": " ", " ": " ", " ": " ", " ": " ",
}
# Bullet glyphs, including Wingdings/Symbol private-use codepoints seen in the corpus.
BULLETS = set("•▪◦‣·§⁃●○■□") | {chr(c) for c in range(0xF000, 0xF0FF)}
# Zero-width / bidi controls. These are also a prompt-injection carrier, so they are
# stripped here and separately reported by sanitize.py.
INVISIBLE = re.compile(r"[​-‏‪-‮⁠-⁤﻿­]")


def clean_text(raw: str) -> tuple[str, list[str]]:
    """-> (cleaned text, list of repairs applied). Repairs are logged, never silent."""
    repairs: list[str] = []
    s = raw

    n_inv = len(INVISIBLE.findall(s))
    if n_inv:
        s = INVISIBLE.sub("", s)
        repairs.append(f"stripped {n_inv} zero-width/bidi control characters")

    lig_hits = sum(s.count(k) for k in LIGATURES if ord(k[0]) > 0x2000 or k in ("Ɵ", "ƞ"))
    for k, v in LIGATURES.items():
        if k in s:
            s = s.replace(k, v)
    if lig_hits:
        repairs.append(f"repaired {lig_hits} ligature/subsetting artefacts (e.g. U+019F->'ti')")

    n_bul = sum(s.count(b) for b in BULLETS)
    if n_bul:
        s = "".join((" • " if c in BULLETS else c) for c in s)
        repairs.append(f"normalised {n_bul} bullet glyphs")

    # Re-join words split across a line by a soft hyphen ("quantita-\ntive").
    s, n_hyp = re.subn(r"(\w)-\n\s*(\w)", r"\1\2", s)
    if n_hyp:
        repairs.append(f"re-joined {n_hyp} hyphen-wrapped words")

    # PDF text layers wrap mid-sentence. Join a line into the next when the break is
    # clearly not a real break: no terminal punctuation and the next line is lowercase.
    s, n_wrap = re.subn(r"([a-z,;])\n(?=[a-z])", r"\1 ", s)
    if n_wrap:
        repairs.append(f"un-wrapped {n_wrap} mid-sentence line breaks")

    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" *\n *", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip(), repairs


# --------------------------------------------------------------------------- results
@dataclass
class Block:
    text: str
    page: int | None
    char_start: int
    char_end: int
    kind: str = "body"          # body | table | header | footer
    column: int = 0
    bbox: tuple[float, float, float, float] | None = None


@dataclass
class Document:
    doc_id: str
    source_file: str
    file_type: str
    text: str
    blocks: list[Block] = field(default_factory=list)
    page_count: int | None = None
    file_sha256: str = ""
    text_sha256: str = ""
    repairs: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    header_footer_text: str = ""
    extraction_quality: float = 0.0
    is_synthetic: bool = False

    def page_of(self, char_pos: int) -> int | None:
        for b in self.blocks:
            if b.char_start <= char_pos < b.char_end:
                return b.page
        return None


def detect_type(path: Path) -> str:
    """Magic bytes, not the file extension. Uploads are untrusted."""
    head = path.read_bytes()[:8]
    if head.startswith(b"%PDF-"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(path) as z:
                names = set(z.namelist())
            if "word/document.xml" in names:
                return "docx"
            return "zip_unsupported"
        except zipfile.BadZipFile:
            return "unknown"
    if head[:5] == b"{\\rtf":
        return "rtf_unsupported"
    return "unknown"


# ----------------------------------------------------------------------------- DOCX
def _para_text(p) -> str:
    """Concatenate runs, but preserve tabs and breaks as whitespace.

    Word CVs use tabs as column separators. Dropping them silently welds the two
    columns together -- 'Ann Arbor, MI' + 'May 2017' becomes 'Ann Arbor, MIMay 2017',
    which then defeats date extraction on that line.
    """
    out: list[str] = []
    for n in p.iter():
        tag = n.tag
        if tag == qn("w:t"):
            out.append(n.text or "")
        elif tag in (qn("w:tab"), qn("w:br"), qn("w:cr")):
            out.append(" ")
    return "".join(out)


def _table_rows(tbl) -> list[list[str]]:
    """Read a table in grid order while collapsing merged cells.

    A <w:tc> that spans N grid columns is returned N times by the OOXML row model;
    keeping identity per row (and skipping vertical-merge continuations) is what
    stops Viktor Sharat's achievements appearing four times each.
    """
    rows: list[list[str]] = []
    for tr in tbl.findall(qn("w:tr")):
        seen: set[int] = set()
        cells: list[str] = []
        for tc in tr.findall(qn("w:tc")):
            if id(tc) in seen:
                continue
            seen.add(id(tc))
            vmerge = tc.find(f"{W_NS}tcPr/{W_NS}vMerge")
            if vmerge is not None and vmerge.get(qn("w:val")) in (None, "continue"):
                continue  # continuation of a vertically merged cell: content is above
            txt = " ".join(_para_text(p) for p in tc.findall(qn("w:p"))).strip()
            cells.append(txt)
        # Collapse a row whose cells are all identical (a fully merged banner row).
        uniq = [c for i, c in enumerate(cells) if c and c not in cells[:i]]
        rows.append(uniq if len(uniq) < len(cells) else cells)
    return rows


def _headers_footers(path: Path) -> tuple[str, list[str]]:
    """Header/footer text, kept apart from the body. Agency watermarks live here."""
    out, notes = [], []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if re.match(r"word/(header|footer)\d*\.xml", name):
                xml = z.read(name).decode("utf8", "ignore")
                txt = " ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)).strip()
                txt = re.sub(r"\s+", " ", txt)
                if txt:
                    out.append(txt)
                    notes.append(f"{name.split('/')[-1]}: {txt[:80]}")
    return "\n".join(dict.fromkeys(out)), notes


def extract_docx(path: Path) -> Document:
    d = docx.Document(str(path))
    body = d.element.body
    parts: list[str] = []
    blocks: list[Block] = []
    pos = 0
    warnings: list[str] = []

    # True document order: iterate body children rather than .paragraphs then .tables.
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            t = _para_text(child).strip()
            if not t:
                continue
            parts.append(t)
            blocks.append(Block(t, None, pos, pos + len(t), "body"))
            pos += len(t) + 1
        elif child.tag == qn("w:tbl"):
            rows = _table_rows(child)
            flat: list[str] = []
            for r in rows:
                cells = [c for c in r if c]
                if not cells:
                    continue
                # A one-value row is prose; a multi-value row is a real record.
                flat.append(cells[0] if len(cells) == 1 else " | ".join(cells))
            if not flat:
                continue
            t = "\n".join(flat)
            parts.append(t)
            blocks.append(Block(t, None, pos, pos + len(t), "table"))
            pos += len(t) + 1

    raw = "\n".join(parts)
    hf_text, hf_notes = _headers_footers(path)
    if hf_notes:
        warnings.extend(f"header/footer content held out of body -> {n}" for n in hf_notes)

    text, repairs = clean_text(raw)
    doc = Document(
        doc_id="", source_file=path.name, file_type="docx", text=text, blocks=blocks,
        page_count=None, repairs=repairs, warnings=warnings, header_footer_text=hf_text,
    )
    _remap_blocks(doc, raw, text)
    return doc


# ------------------------------------------------------------------------------ PDF
def _columns(blocks: list[tuple], page_width: float,
             page_height: float) -> tuple[list[int], str]:
    """Assign each block a column index, but only when the page is genuinely multi-column.

    An x-gap alone is not evidence of a column. A centred name block sits far to the
    right of the body text and produces exactly the same gap -- and treating it as a
    column moves the candidate's name to the bottom of the document, which is how this
    detector originally broke `RYAN PATEL - Resume.pdf` (a single-column CV) while
    fixing `Omar El-Hassan 202405.pdf` (a real two-column one).

    A real column therefore has to look like a column: enough blocks to be a body of
    text, spanning enough of the page vertically, and running *alongside* its neighbour
    rather than sitting above it. On Ryan's page the right-hand group is 2 blocks in a
    narrow band near the top and is correctly rejected; on Omar's it is 15 blocks
    spanning the page and is correctly accepted.
    """
    if len(blocks) < 4:
        return [0] * len(blocks), "single column (too few blocks to be multi-column)"

    xs = sorted({round(b[0]) for b in blocks})
    gap = max(70.0, page_width * 0.12)
    candidates = [(a + b) / 2 for a, b in zip(xs, xs[1:]) if b - a > gap]

    MIN_BLOCKS = 3                       # a column is a body of text, not a stray label
    MIN_V_SPAN = 0.30 * page_height      # ...and it runs down a real part of the page
    MIN_OVERLAP = 0.35                   # ...beside its neighbour, not above it

    accepted: list[float] = []
    for bound in candidates:
        left = [b for b in blocks if b[0] <= bound]
        right = [b for b in blocks if b[0] > bound]
        if len(left) < MIN_BLOCKS or len(right) < MIN_BLOCKS:
            continue
        l0, l1 = min(b[1] for b in left), max(b[3] for b in left)
        r0, r1 = min(b[1] for b in right), max(b[3] for b in right)
        if (l1 - l0) < MIN_V_SPAN or (r1 - r0) < MIN_V_SPAN:
            continue
        overlap = min(l1, r1) - max(l0, r0)
        if overlap < MIN_OVERLAP * min(l1 - l0, r1 - r0):
            continue
        accepted.append(bound)

    if not accepted:
        why = ("single column" if not candidates else
               "single column (x-gap present but it is a centred header or a stray "
               "label, not a column: too few blocks, too little vertical span, or no "
               "side-by-side overlap)")
        return [0] * len(blocks), why

    def col_of(x0: float) -> int:
        return sum(1 for bd in accepted if x0 > bd)

    return [col_of(b[0]) for b in blocks], f"{len(accepted) + 1}-column layout"


def extract_pdf(path: Path) -> Document:
    import fitz  # imported lazily: keeps notebook import time down

    d = fitz.open(str(path))
    parts: list[str] = []
    blocks_out: list[Block] = []
    pos = 0
    warnings: list[str] = []

    for pno, page in enumerate(d, start=1):
        raw_blocks = [b for b in page.get_text("blocks") if b[6] == 0 and b[4].strip()]
        cols, why = _columns(raw_blocks, page.rect.width, page.rect.height)
        n_cols = len(set(cols))
        if n_cols > 1:
            warnings.append(
                f"page {pno}: detected {why}; reading order repaired column-major "
                f"(naive extraction interleaves the columns)")
        # Column-major, then top-to-bottom, then left-to-right within a band.
        order = sorted(range(len(raw_blocks)),
                       key=lambda i: (cols[i], round(raw_blocks[i][1] / 6), raw_blocks[i][0]))
        for i in order:
            b = raw_blocks[i]
            t = re.sub(r"\n+", "\n", b[4]).strip()
            if not t:
                continue
            parts.append(t)
            blocks_out.append(Block(t, pno, pos, pos + len(t), "body", cols[i],
                                    (b[0], b[1], b[2], b[3])))
            pos += len(t) + 1

    if not "".join(parts).strip():
        warnings.append("no text layer found -- this document requires OCR (flagged, "
                        "OCR is off by default)")

    raw = "\n".join(parts)
    text, repairs = clean_text(raw)
    doc = Document(doc_id="", source_file=path.name, file_type="pdf", text=text,
                   blocks=blocks_out, page_count=d.page_count, repairs=repairs,
                   warnings=warnings)
    d.close()
    _remap_blocks(doc, raw, text)
    return doc


# ------------------------------------------------------------------------- assembly
def _remap_blocks(doc: Document, raw: str, cleaned: str) -> None:
    """Block offsets were computed against raw text; cleaning shifted them.

    Rather than track edits, re-locate each block's cleaned form in the cleaned text
    with a forward-only cursor. Forward-only matters: it prevents a repeated line (a
    duplicated table banner) from stealing an earlier block's offsets.
    """
    cursor = 0
    for b in doc.blocks:
        needle, _ = clean_text(b.text)
        needle = needle.split("\n")[0][:60]
        if not needle:
            continue
        idx = cleaned.find(needle, cursor)
        if idx == -1:
            idx = cleaned.find(needle)
        if idx == -1:
            continue
        b.char_start = idx
        b.char_end = min(len(cleaned), idx + len(b.text))
        cursor = idx


_ALNUM = re.compile(r"[A-Za-z0-9]")


def score_extraction_quality(doc: Document) -> float:
    """0-1 heuristic used to route documents to review and to weight confidence."""
    t = doc.text
    if len(t) < 200:
        return 0.05
    alnum = len(_ALNUM.findall(t)) / max(1, len(t))
    # Ratio of recognisable words; mojibake and OCR noise crater this.
    words = re.findall(r"[A-Za-z]{2,}", t)
    vowelly = sum(1 for w in words if re.search(r"[aeiouAEIOU]", w)) / max(1, len(words))
    weird = len(re.findall(r"[^\x00-\x7F]", t)) / max(1, len(t))
    lines = [l for l in t.split("\n") if l.strip()]
    avg_len = sum(len(l) for l in lines) / max(1, len(lines))
    length_ok = min(1.0, avg_len / 40)
    score = 0.35 * alnum + 0.30 * vowelly + 0.20 * length_ok + 0.15 * (1 - min(1.0, weird * 20))
    return round(max(0.0, min(1.0, score)), 3)


def load_document(path: Path, is_synthetic: bool = False) -> Document:
    path = Path(path)
    file_sha = hashlib.sha256(path.read_bytes()).hexdigest()
    ftype = detect_type(path)
    if ftype == "pdf":
        doc = extract_pdf(path)
    elif ftype == "docx":
        doc = extract_docx(path)
    else:
        doc = Document(doc_id="", source_file=path.name, file_type=ftype, text="",
                       warnings=[f"unsupported file type: {ftype}"])
    doc.file_sha256 = file_sha
    doc.text_sha256 = hashlib.sha256(doc.text.encode()).hexdigest()
    doc.doc_id = file_sha[:16]
    doc.is_synthetic = is_synthetic
    doc.extraction_quality = score_extraction_quality(doc)
    if doc.extraction_quality < 0.55:
        doc.warnings.append(
            f"low extraction quality ({doc.extraction_quality}) -- downstream confidence "
            f"is capped and the document is routed to human review")
    return doc
