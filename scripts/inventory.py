"""Raw data inventory over the 10 supplied resumes. No design decisions before this runs."""
import pathlib, re, zipfile, hashlib, json, sys
import fitz  # PyMuPDF
import docx

ROOT = pathlib.Path(__file__).resolve().parent.parent
FILES = sorted([p for p in ROOT.iterdir() if p.suffix.lower() in {".pdf", ".docx"}])

def docx_text(p):
    d = docx.Document(str(p))
    parts = [pa.text for pa in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts), len(d.tables), len(d.paragraphs)

def docx_raw_xml_extras(p):
    """Text living in headers/footers/textboxes that python-docx paragraphs miss."""
    extras = []
    with zipfile.ZipFile(p) as z:
        for n in z.namelist():
            if n.startswith("word/") and n.endswith(".xml") and ("header" in n or "footer" in n):
                xml = z.read(n).decode("utf8", "ignore")
                txt = re.sub(r"<[^>]+>", " ", xml)
                txt = re.sub(r"\s+", " ", txt).strip()
                if txt: extras.append((n, txt[:200]))
    return extras

def pdf_text(p):
    doc = fitz.open(str(p))
    pages = []
    for pg in doc:
        pages.append(pg.get_text("text"))
    return "\n".join(pages), doc.page_count

SEC = re.compile(r"^(?P<h>[A-Z][A-Z &/\-\.']{3,60})\s*$", re.M)
EMAIL = re.compile(r"[\w\.\-+]+@[\w\-]+\.\w+")
PHONE = re.compile(r"(\+?\d[\d\s\-\(\)\.]{7,}\d)")
DATE = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\b(19|20)\d{2}\b")
NONASCII = re.compile(r"[^\x00-\x7F]")

rows = []
for p in FILES:
    rec = {"file": p.name, "ext": p.suffix.lower(), "bytes": p.stat().st_size}
    if p.suffix.lower() == ".docx":
        text, ntab, npar = docx_text(p)
        rec["pages"] = None; rec["tables"] = ntab; rec["paragraphs"] = npar
        rec["hdr_ftr"] = len(docx_raw_xml_extras(p))
    else:
        text, npg = pdf_text(p)
        rec["pages"] = npg; rec["tables"] = None; rec["paragraphs"] = None
        rec["hdr_ftr"] = None
    rec["chars"] = len(text)
    rec["lines"] = text.count("\n") + 1
    rec["has_text_layer"] = len(text.strip()) > 200
    rec["emails"] = EMAIL.findall(text)
    rec["phones"] = [x.strip() for x in PHONE.findall(text)][:3]
    rec["n_dates"] = len(DATE.findall(text))
    rec["nonascii"] = sorted(set(NONASCII.findall(text)))[:12]
    rec["headings"] = [m.group("h").strip() for m in SEC.finditer(text)][:20]
    rec["sha"] = hashlib.sha256(text.encode()).hexdigest()[:12]
    rec["_text"] = text
    rows.append(rec)

out = ROOT / "data" / "raw_text"
out.mkdir(parents=True, exist_ok=True)
for r in rows:
    (out / (pathlib.Path(r["file"]).stem + ".txt")).write_text(r["_text"])

print(f"{'FILE':42s} {'EXT':5s} {'PG':>3s} {'CHARS':>6s} {'LINES':>5s} {'TBL':>3s} {'DATES':>5s} {'TEXT?':>5s}")
for r in rows:
    print(f"{r['file'][:42]:42s} {r['ext']:5s} {str(r['pages'] or '-'):>3s} {r['chars']:>6d} {r['lines']:>5d} {str(r['tables'] if r['tables'] is not None else '-'):>3s} {r['n_dates']:>5d} {str(r['has_text_layer']):>5s}")

print("\n=== CONTACT / UNICODE / HEADINGS ===")
for r in rows:
    print(f"\n--- {r['file']}")
    print("  emails :", r["emails"])
    print("  phones :", r["phones"])
    print("  nonascii:", r["nonascii"])
    print("  headings:", r["headings"])
